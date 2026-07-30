from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = OUTPUT_DIR / "Shahil_Shaik_Academic_CV.docx"

# Design system:
# - Base preset: compact_reference_guide.
# - First-page pattern: proposal_centerpiece, adapted into a compact identity masthead.
# - Named academic-CV override: 10.25 pt body, 1.08 line spacing, and restrained
#   section spacing to keep the application CV within two pages.
FONT = "Calibri"
NAVY = RGBColor(24, 55, 86)
BLUE = RGBColor(46, 116, 181)
INK = RGBColor(28, 33, 38)
MUTED = RGBColor(91, 101, 111)
LIGHT = RGBColor(218, 226, 234)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run(run, size=10.25, color=INK, bold=False, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True, size=9.5):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    run_props.append(fonts)
    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), str(color))
    run_props.append(color_node)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    run_props.append(size_node)
    if underline:
        underline_node = OxmlElement("w:u")
        underline_node.set(qn("w:val"), "single")
        run_props.append(underline_node)
    run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_paragraph_border(paragraph, color="DADFE5", size="6", space="2"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run(run, size=8.5, color=MUTED)


def add_section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title.upper())
    set_run(r, size=10.5, color=NAVY, bold=True)
    set_paragraph_border(p)
    return p


def add_body(doc, text="", after=3, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.keep_together = keep
    set_run(p.add_run(text))
    return p


def add_project(doc, title, subtitle, description):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.06
    p.paragraph_format.keep_together = True
    set_run(p.add_run(title), bold=True, color=NAVY)
    set_run(p.add_run(f"  |  {subtitle}"), size=9.6, color=MUTED, italic=True)
    set_run(p.add_run(f"\n{description}"))
    return p


def add_publication(doc, authors, year, title, venue):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.keep_together = True
    before, match, after = authors.partition("Shaik, S.")
    if before:
        set_run(p.add_run(before))
    if match:
        set_run(p.add_run(match), bold=True, color=NAVY)
    if after:
        set_run(p.add_run(after))
    set_run(p.add_run(f" ({year}). "))
    set_run(p.add_run(title), italic=True)
    set_run(p.add_run(f". {venue}"))
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.82)
section.right_margin = Inches(0.82)
section.header_distance = Inches(0.32)
section.footer_distance = Inches(0.32)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(10.25)
normal.font.color.rgb = INK
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.08

for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, NAVY, 10, 5),
):
    style = styles[style_name]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header
header_p = header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header_p.paragraph_format.space_after = Pt(0)
set_run(header_p.add_run("SHAHIL SHAIK  |  ACADEMIC CV"), size=8, color=MUTED, bold=True)

footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(0)
footer_p.paragraph_format.space_after = Pt(0)
set_run(footer_p.add_run("Shahil Shaik  •  Page "), size=8.5, color=MUTED)
add_page_field(footer_p)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(1)
title.paragraph_format.space_after = Pt(1)
set_run(title.add_run("SHAHIL SHAIK"), size=22, color=NAVY, bold=True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(3)
set_run(
    subtitle.add_run("Ph.D. Researcher in Robotics and Intelligent Systems"),
    size=11.5,
    color=MUTED,
    bold=True,
)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(5)
set_run(contact.add_run("Clemson University  •  "), size=9.5, color=MUTED)
add_hyperlink(contact, "shahils@clemson.edu", "mailto:shahils@clemson.edu")
set_run(contact.add_run("  •  "), size=9.5, color=MUTED)
add_hyperlink(
    contact,
    "Google Scholar",
    "https://scholar.google.com/citations?user=THKUDwwAAAAJ&hl=en",
)
set_run(contact.add_run("  •  "), size=9.5, color=MUTED)
add_hyperlink(contact, "GitHub", "https://github.com/shahils01")

add_section_heading(doc, "Research Profile")
p = add_body(doc, keep=True)
set_run(
    p.add_run(
        "Robotics Ph.D. researcher developing efficient learning architectures for "
        "embodied and multi-agent intelligence. Research spans transformer attention, "
        "vision-language and vision-language-action models, multi-agent deep reinforcement "
        "learning, distributional policy optimization, and shared human–robot control. "
        "Published in the American Control Conference and IET Cyber-Systems and Robotics; "
        "experienced in reproducible, GPU-accelerated research with PyTorch."
    )
)

add_section_heading(doc, "Education")
education = doc.add_table(rows=3, cols=2)
education.autofit = False
education.columns[0].width = Inches(1.25)
education.columns[1].width = Inches(5.45)
tbl = education._tbl
tbl_pr = tbl.tblPr
tbl_layout = OxmlElement("w:tblLayout")
tbl_layout.set(qn("w:type"), "fixed")
tbl_pr.append(tbl_layout)
tbl_width = OxmlElement("w:tblW")
tbl_width.set(qn("w:w"), "9648")
tbl_width.set(qn("w:type"), "dxa")
tbl_pr.append(tbl_width)
for row, (dates, degree, school, detail) in zip(
    education.rows,
    (
        (
            "2021–Present",
            "Ph.D. in Mechanical Engineering (Robotics)",
            "Clemson University, Clemson, South Carolina",
            "Research focus: intelligent multi-agent robotic systems and learning-based autonomy.",
        ),
        (
            "2019–2021",
            "Graduate Study in Mechanical Engineering",
            "Clemson University, Clemson, South Carolina",
            "Entered the M.S. program in 2019 and transitioned to the Ph.D. program in 2021.",
        ),
        (
            "2015–2019",
            "Bachelor’s Degree in Mechanical Engineering",
            "Anil Neerukonda Institute of Technology and Sciences, Visakhapatnam, India",
            "GPA: 8.55/10.",
        ),
    ),
):
    for c in row.cells:
        set_cell_margins(c, top=15, start=0, bottom=30, end=80)
        c.width = Inches(1.25 if c is row.cells[0] else 5.45)
    left = row.cells[0].paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    set_run(left.add_run(dates), size=9.5, color=MUTED, bold=True)
    right = row.cells[1].paragraphs[0]
    right.paragraph_format.space_after = Pt(0)
    right.paragraph_format.line_spacing = 1.03
    set_run(right.add_run(degree), bold=True, color=NAVY)
    set_run(right.add_run(f"\n{school}"), size=9.8)
    set_run(right.add_run(f"\n{detail}"), size=9.4, color=MUTED)

add_section_heading(doc, "Research Interests")
add_body(
    doc,
    "Efficient transformer architectures  •  Vision-language-action learning  •  "
    "Multi-agent reinforcement learning  •  Robot learning  •  Human–robot shared control  •  "
    "Distributed and memory-efficient deep learning",
    keep=True,
)

add_section_heading(doc, "Selected Research")
add_project(
    doc,
    "Lie-Generated Metric Attention (LGMA)",
    "Lead developer and investigator",
    "Developing a PyTorch attention architecture that generates head-specific metrics from "
    "shared Lie-algebra generators to reduce parameter and key–value cache costs. Built "
    "controlled MHA-versus-LGMA experiments, stability diagnostics, checkpointing, efficiency "
    "accounting, and single-/multi-GPU training workflows for TinyStories and robot-learning studies.",
)
add_project(
    doc,
    "Vision-Language(-Action) Models for Multi-Agent Robotics",
    "Research and software development",
    "Investigating multimodal critics and action models for policy evaluation and robot learning, "
    "including MA-VLCM and planned LIBERO benchmarking. Work integrates visual observations, "
    "language-conditioned reasoning, temporal context, and multi-agent state representations.",
)
add_project(
    doc,
    "Multi-Agent Deep Reinforcement Learning",
    "Doctoral research",
    "Developed methods for learning under constrained inter-agent communication and studied "
    "distributional policy-gradient estimation. Research emphasizes coordination, robustness, "
    "sample efficiency, and principled evaluation across multi-agent settings.",
)
add_project(
    doc,
    "Shared Control of Mobile Robots",
    "Collaborative research",
    "Studied deep reinforcement-learning approaches that combine human guidance with autonomous "
    "control for mobile robot navigation, resulting in a peer-reviewed journal publication.",
)

doc.add_page_break()
add_section_heading(doc, "Publications and Preprints")
add_publication(
    doc,
    "Shaik, S., Smereka, J. M., & Wang, Y.",
    "2026",
    "Multi-Agent Deep Reinforcement Learning Under Constrained Communications",
    "arXiv preprint arXiv:2601.17069.",
)
add_publication(
    doc,
    "Shaik, S., Parameshwaran, A., Nayak, A., Smereka, J. M., & Wang, Y.",
    "2026",
    "MA-VLCM: A Vision Language Critic Model for Value Estimation of Policies in Multi-Agent Team Settings",
    "arXiv preprint arXiv:2603.15418.",
)
add_publication(
    doc,
    "Nayak, A., Shaik, S., & Wang, Y.",
    "2026",
    "Belief-Aware VLM Model for Human-like Reasoning",
    "arXiv preprint arXiv:2604.09686.",
)
add_publication(
    doc,
    "Shaik, S., Smereka, J. M., & Wang, Y.",
    "2025",
    "Generalized Advantage Estimation for Distributional Policy Gradients",
    "Proceedings of the American Control Conference, 3073–3078.",
)
add_publication(
    doc,
    "Tian, C., Shaik, S., & Wang, Y.",
    "2021",
    "Deep Reinforcement Learning for Shared Control of Mobile Robots",
    "IET Cyber-Systems and Robotics, 3(4), 315–330.",
)

add_section_heading(doc, "Research Software and Reproducibility")
software = add_body(doc, after=3, keep=True)
set_run(software.add_run("Open-source LGMA implementation. "), bold=True, color=NAVY)
set_run(
    software.add_run(
        "Maintains a public PyTorch research prototype with interchangeable LGMA and conventional "
        "multi-head attention, synthetic and TinyStories experiments, and VLA data/model/evaluation modules. "
    )
)
add_hyperlink(
    software,
    "Repository",
    "https://github.com/shahils01/Attention-is-Transformed",
    size=10.25,
)

repro = add_body(doc, after=3, keep=True)
set_run(repro.add_run("Experiment integrity. "), bold=True, color=NAVY)
set_run(
    repro.add_run(
        "Built unit and smoke tests, configuration-driven experiments, resumable checkpoints, "
        "validation workflows, per-GPU throughput and memory accounting, and automated summaries "
        "for validation-loss-versus-token and validation-loss-versus-GPU-hour comparisons."
    )
)

hpc = add_body(doc, after=3, keep=True)
set_run(hpc.add_run("Research-computing readiness. "), bold=True, color=NAVY)
set_run(
    hpc.add_run(
        "Prepared bf16 CUDA training, gradient accumulation, single- and multi-GPU "
        "DistributedDataParallel execution, Slurm launch patterns, dataset staging, and "
        "portable Conda/Apptainer environment plans for accelerator-based studies."
    )
)

add_section_heading(doc, "Technical Skills")
skills = doc.add_table(rows=5, cols=2)
skills.autofit = False
skills.columns[0].width = Inches(1.48)
skills.columns[1].width = Inches(5.22)
skills_tbl_pr = skills._tbl.tblPr
skills_layout = OxmlElement("w:tblLayout")
skills_layout.set(qn("w:type"), "fixed")
skills_tbl_pr.append(skills_layout)
skills_width = OxmlElement("w:tblW")
skills_width.set(qn("w:w"), "9648")
skills_width.set(qn("w:type"), "dxa")
skills_tbl_pr.append(skills_width)
for row, (label, value) in zip(
    skills.rows,
    (
        ("Programming", "Python; scientific computing and research software development"),
        ("Deep Learning", "PyTorch; transformer models; attention mechanisms; reinforcement learning; VLM/VLA"),
        ("GPU & Scale", "CUDA; bf16 mixed precision; PyTorch DistributedDataParallel; torchrun; Slurm"),
        ("Research Stack", "Hugging Face Datasets/Transformers; NumPy; h5py; OpenCV; Weights & Biases"),
        ("Robotics & Tools", "MuJoCo/LIBERO workflows; Git/GitHub; Conda; Apptainer; pytest"),
    ),
):
    for c in row.cells:
        set_cell_margins(c, top=18, start=0, bottom=18, end=80)
    lp = row.cells[0].paragraphs[0]
    lp.paragraph_format.space_after = Pt(0)
    set_run(lp.add_run(label), size=9.5, color=NAVY, bold=True)
    rp = row.cells[1].paragraphs[0]
    rp.paragraph_format.space_after = Pt(0)
    set_run(rp.add_run(value), size=9.5)

doc.core_properties.title = "Shahil Shaik — Academic CV"
doc.core_properties.subject = "Academic CV for research computing allocation applications"
doc.core_properties.author = "Shahil Shaik"
doc.core_properties.keywords = "robotics, reinforcement learning, transformers, VLA, multi-agent systems"
doc.core_properties.comments = ""

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
