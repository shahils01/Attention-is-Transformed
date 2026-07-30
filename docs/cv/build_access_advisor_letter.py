from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
OUT_PATH = OUT_DIR / "ACCESS_Advisor_Support_Letter_Draft.docx"

# Formal academic support letter.
# Base preset: standard_business_brief.
# Named override: 1.2-inch top margin leaves space for Clemson letterhead.
FONT = "Calibri"
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(90, 98, 106)


def style_run(run, *, size=11, bold=False, italic=False, color=BLACK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run


def add_para(doc, text="", *, after=8, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_together = keep
    style_run(p.add_run(text))
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1.20)
section.bottom_margin = Inches(0.85)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(11)
normal.font.color.rgb = BLACK
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.10

# The top margin is intentionally open so the advisor can transfer the text to
# official Clemson University letterhead without disturbing the body layout.
date_p = add_para(doc, "July 27, 2026", after=12)
date_p.paragraph_format.keep_with_next = True

for line, after in (
    ("ACCESS Allocations Review Committee", 0),
    ("Advanced Cyberinfrastructure Coordination Ecosystem: Services & Support (ACCESS)", 0),
    ("Re: Advisor Support for Shahil Shaik’s Explore ACCESS Request", 12),
):
    p = add_para(doc, after=after)
    p.paragraph_format.keep_with_next = True
    style_run(p.add_run(line), bold=line.startswith("Re:"))

add_para(doc, "Dear ACCESS Allocations Review Committee:", after=10)

p = add_para(doc)
style_run(p.add_run("I am writing in support of "))
style_run(p.add_run("Shahil Shaik’s Explore ACCESS allocation request"), bold=True)
style_run(
    p.add_run(
        " for NCSA DeltaAI. Shahil is a Ph.D. Candidate and Graduate Research Assistant "
        "in the Department of Mechanical Engineering at Clemson University, and I serve as "
        "his faculty advisor."
    )
)

p = add_para(doc)
style_run(p.add_run("I am aware of and support the proposed project, "))
style_run(
    p.add_run(
        "“Evaluating Lie-Generated Metric Attention for Efficient Language and "
        "Vision-Language-Action Models.”"
    ),
    italic=True,
)
style_run(
    p.add_run(
        " The research will evaluate Lie-Generated Metric Attention as a parameter- and "
        "memory-efficient alternative to conventional multi-head attention. The planned work "
        "includes controlled language-model experiments using TinyStories and, where appropriate, "
        "a WikiHow corpus, together with vision-language-action experiments on the LIBERO "
        "robot-learning benchmark."
    )
)

add_para(
    doc,
    "NCSA DeltaAI is appropriate for this work because the experiments require H100/GH200-class "
    "GPU memory and mixed-precision throughput for repeated transformer training, multimodal "
    "vision processing, and selected multi-GPU scaling studies. The project will compare model "
    "quality, convergence, parameter efficiency, key-value cache requirements, throughput, peak "
    "GPU memory, and task performance using reproducible experimental configurations.",
)

add_para(
    doc,
    "I intend to engage with and guide the computational and research activities conducted under "
    "this allocation. My guidance will include oversight of experimental design, responsible use "
    "of the requested resources, interpretation of results, reproducibility, and dissemination of "
    "the resulting research. I will remain involved throughout the allocation period and will "
    "support Shahil in meeting applicable ACCESS reporting and resource-use requirements.",
)

add_para(
    doc,
    "I believe the requested resources will materially advance this research and provide an "
    "appropriate computational environment for rigorous evaluation of the proposed attention and "
    "robot-learning methods. I support Shahil’s request for ACCESS resources.",
    after=14,
)

add_para(doc, "Sincerely,", after=34)

signature = add_para(doc, after=0, keep=True)
style_run(signature.add_run("Yue Wang, Ph.D."), bold=True)
add_para(doc, "Faculty Advisor", after=0, keep=True)
add_para(doc, "Department of Mechanical Engineering", after=0, keep=True)
add_para(doc, "Clemson University", after=0, keep=True)

footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(0)
footer_p.paragraph_format.space_after = Pt(0)
style_run(
    footer_p.add_run("Draft for review — transfer to official Clemson University letterhead before signature"),
    size=8,
    italic=True,
    color=MUTED,
)

doc.core_properties.title = "Advisor Support Letter for Shahil Shaik’s Explore ACCESS Request"
doc.core_properties.subject = "NCSA DeltaAI allocation request"
doc.core_properties.author = "Yue Wang"
doc.core_properties.comments = ""

OUT_DIR.mkdir(parents=True, exist_ok=True)
doc.save(OUT_PATH)
print(OUT_PATH)
